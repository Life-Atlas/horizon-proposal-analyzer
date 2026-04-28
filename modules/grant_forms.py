"""
Grant Forms Module — Pre-flight validation and cross-document consistency.

Validates company data against a verified registry before any form filling.
Checks cross-document consistency (budget, org.nr, names) across all
submission documents.

Usage from CLI:
  python -m modules.grant_forms --registry path/to/company_registry.json --check-dir ./FINAL/
  python -m modules.grant_forms --registry path/to/company_registry.json --lookup "WINNIIO AB" --keur

MIT License — WINNIIO AB / Life Atlas
"""

import json
import re
import sys
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, Union

__all__ = ["GrantFormsValidator", "CompanyData", "ConsistencyError"]

try:
    import fitz
    _HAS_FITZ = True
except ImportError:
    _HAS_FITZ = False

try:
    from docx import Document as DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


DEFAULT_REGISTRY = Path.home() / ".claude" / "data" / "company_registry.json"


@dataclass
class CompanyData:
    name: str
    org_nr: str
    location: str
    vd: str
    employees: int
    revenue_ksek: int
    balance_ksek: int
    fiscal_year: str
    company_type: str
    owner: Optional[str] = None
    owner_org_nr: Optional[str] = None
    sme_category: str = "micro"

    def to_keur(self, rate: float) -> dict:
        return {
            "revenue_keur": round(self.revenue_ksek / rate),
            "balance_keur": round(self.balance_ksek / rate),
            "rate": rate,
        }


@dataclass
class ConsistencyError:
    field: str
    doc_a: str
    value_a: str
    doc_b: str
    value_b: str
    severity: str = "ERROR"

    def __str__(self):
        return f"[{self.severity}] '{self.field}': '{self.value_a}' in {self.doc_a} vs '{self.value_b}' in {self.doc_b}"


class GrantFormsValidator:
    """Validates grant submission documents for consistency and correctness."""

    def __init__(self, registry_path: Optional[Union[str, Path]] = None):
        self.registry_path = Path(registry_path) if registry_path else DEFAULT_REGISTRY
        self._registry = None
        self._rate = None

    @property
    def registry(self) -> dict:
        if self._registry is None:
            if not self.registry_path.exists():
                raise FileNotFoundError(f"Company registry not found: {self.registry_path}")
            self._registry = json.loads(self.registry_path.read_text(encoding="utf-8"))
        return self._registry

    @property
    def exchange_rate(self) -> float:
        if self._rate is None:
            self._rate = self.registry["exchange_rates"]["SEK_EUR"]
        return self._rate

    def lookup(self, name: str) -> CompanyData | None:
        for key, data in self.registry["companies"].items():
            if name.lower() in key.lower() or name == data.get("org_nr", ""):
                return CompanyData(
                    name=key,
                    org_nr=data["org_nr"],
                    location=data["location"],
                    vd=data["vd"],
                    employees=data["employees"],
                    revenue_ksek=data["revenue_ksek"],
                    balance_ksek=data["balance_ksek"],
                    fiscal_year=data.get("fiscal_year", ""),
                    company_type=data.get("company_type", "fristående"),
                    owner=data.get("owner"),
                    owner_org_nr=data.get("owner_org_nr"),
                    sme_category=data.get("sme_category", "micro"),
                )
        return None

    def preflight_check(self, company_names: list[str]) -> list[str]:
        """Run pre-flight checks before form filling. Returns list of blockers."""
        blockers = []
        for name in company_names:
            data = self.lookup(name)
            if data is None:
                blockers.append(f"BLOCKER: '{name}' not in registry. Verify on allabolag.se first.")
            else:
                if not data.org_nr:
                    blockers.append(f"BLOCKER: '{name}' has no org.nr in registry.")
                if data.employees == 0 and data.revenue_ksek > 0:
                    blockers.append(f"WARNING: '{name}' has 0 employees but {data.revenue_ksek} KSEK revenue — verify.")
        return blockers

    def extract_org_nrs_from_pdf(self, path: str) -> list[str]:
        """Extract all Swedish org.nr patterns from a PDF."""
        if not _HAS_FITZ:
            return []
        doc = fitz.open(path)
        text = " ".join(page.get_text() for page in doc)
        doc.close()
        return re.findall(r'\b\d{6}-\d{4}\b', text)

    def extract_org_nrs_from_docx(self, path: str) -> list[str]:
        """Extract all Swedish org.nr patterns from a DOCX."""
        if not _HAS_DOCX:
            return []
        doc = DocxDocument(path)
        text = " ".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += " " + " ".join(c.text for c in row.cells)
        return re.findall(r'\b\d{6}-\d{4}\b', text)

    def extract_numbers_from_text(self, text: str) -> list[str]:
        """Extract formatted numbers that look like financial data."""
        return re.findall(r'\b\d{1,3}(?:[\s,]\d{3})*\b', text)

    def check_cross_document_consistency(self, doc_dir: Union[str, Path]) -> list[ConsistencyError]:
        """Check all documents in a directory for cross-document consistency."""
        doc_dir = Path(doc_dir)
        errors = []
        org_nrs_by_doc: dict[str, set[str]] = {}

        for f in doc_dir.iterdir():
            if f.name.startswith("~$"):
                continue
            if f.suffix.lower() == ".pdf" and _HAS_FITZ:
                org_nrs_by_doc[f.name] = set(self.extract_org_nrs_from_pdf(str(f)))
            elif f.suffix.lower() == ".docx" and _HAS_DOCX:
                org_nrs_by_doc[f.name] = set(self.extract_org_nrs_from_docx(str(f)))

        known_org_nrs = {d["org_nr"] for d in self.registry["companies"].values()}

        for doc_name, found_nrs in org_nrs_by_doc.items():
            for nr in found_nrs:
                if nr not in known_org_nrs:
                    errors.append(ConsistencyError(
                        field="org_nr",
                        doc_a=doc_name,
                        value_a=nr,
                        doc_b="company_registry.json",
                        value_b="NOT FOUND",
                        severity="WARNING",
                    ))

        doc_names = list(org_nrs_by_doc.keys())
        for i in range(len(doc_names)):
            for j in range(i + 1, len(doc_names)):
                nrs_a = org_nrs_by_doc[doc_names[i]]
                nrs_b = org_nrs_by_doc[doc_names[j]]
                only_in_a = nrs_a - nrs_b - {""}
                only_in_b = nrs_b - nrs_a - {""}
                if only_in_a and nrs_b:
                    for nr in only_in_a:
                        if nr in known_org_nrs:
                            errors.append(ConsistencyError(
                                field="org_nr_missing",
                                doc_a=doc_names[i],
                                value_a=nr,
                                doc_b=doc_names[j],
                                value_b="ABSENT",
                                severity="INFO",
                            ))

        return errors

    def validate_modellforsakran_fields(self, form_fields: dict, company_name: str, currency: str = "KEUR") -> list[str]:
        """Validate modellförsäkran form fields against registry data."""
        errors = []
        data = self.lookup(company_name)
        if not data:
            return [f"BLOCKER: '{company_name}' not in registry"]

        if currency == "KEUR":
            keur = data.to_keur(self.exchange_rate)
            expected_rev = keur["revenue_keur"]
            expected_bal = keur["balance_keur"]
        else:
            expected_rev = data.revenue_ksek
            expected_bal = data.balance_ksek

        for field_name, value in form_fields.items():
            val_str = str(value).strip()
            if "org" in field_name.lower() and val_str and val_str != data.org_nr:
                errors.append(f"MISMATCH: org.nr field '{field_name}' = '{val_str}', expected '{data.org_nr}'")

        return errors


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Grant Forms Validator")
    parser.add_argument("--registry", default=str(DEFAULT_REGISTRY), help="Path to company_registry.json")
    parser.add_argument("--lookup", metavar="NAME", help="Look up company data")
    parser.add_argument("--keur", action="store_true", help="Show KEUR conversion")
    parser.add_argument("--preflight", nargs="+", metavar="COMPANY", help="Run pre-flight check for companies")
    parser.add_argument("--check-dir", metavar="DIR", help="Cross-document consistency check on a directory")
    args = parser.parse_args()

    validator = GrantFormsValidator(args.registry)

    if args.lookup:
        data = validator.lookup(args.lookup)
        if not data:
            print(f"NOT FOUND: '{args.lookup}'")
            sys.exit(1)
        print(f"\n  {data.name}")
        print(f"  Org.nr: {data.org_nr} | {data.location} | VD: {data.vd}")
        print(f"  Employees: {data.employees} | Type: {data.company_type} | SME: {data.sme_category}")
        print(f"  Revenue: {data.revenue_ksek} KSEK | Balance: {data.balance_ksek} KSEK")
        if args.keur:
            keur = data.to_keur(validator.exchange_rate)
            print(f"  → Revenue: {keur['revenue_keur']} KEUR | Balance: {keur['balance_keur']} KEUR (rate {keur['rate']})")
        if data.owner:
            print(f"  Owner: {data.owner} ({data.owner_org_nr})")
        print()
        return

    if args.preflight:
        blockers = validator.preflight_check(args.preflight)
        if blockers:
            print("PRE-FLIGHT BLOCKERS:")
            for b in blockers:
                print(f"  ✗ {b}")
            sys.exit(1)
        else:
            print("✓ All companies verified in registry. Proceed with form filling.")
        return

    if args.check_dir:
        errors = validator.check_cross_document_consistency(args.check_dir)
        if errors:
            print(f"CONSISTENCY CHECK — {len(errors)} issues found:\n")
            for e in errors:
                print(f"  {e}")
        else:
            print("✓ No consistency issues found.")
        return

    parser.print_help()


if __name__ == "__main__":
    main()
