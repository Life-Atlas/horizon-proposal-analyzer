#!/usr/bin/env python3
"""
C.R.U.C.I.B.L.E. Light v1.1.0
Call-specific Quick-Score Engine — Frontloaded Template Exploration

LIGHT vs DEEP:
  Light — fast, call-specific, 8 dimensions, iterative, for grant WRITING.
          Designed to be run repeatedly as the author adds content.
          Each pass shows delta from previous, stopping when target is met.
  Deep  — comprehensive, 7 passes, 48+ detectors, for grant REVIEWING.
          Full structural/SMILE/PESTLE analysis in crucible.py.

Architecture:
  Phase 0: TEMPLATE EXPLORATION — unpack .docx, map every field/placeholder/section
  Phase 1: CALL MAPPING — map template sections to call evaluation criteria
  Phase 2: CONTENT SCORING — score each dimension 1-10 against call text
  Phase 3: GAP ANALYSIS — identify what's missing or weak
  Phase 4: ITERATE — fix gaps, re-score, repeat until target

Scoring target: 10.0 across all dimensions.
  - After 3 consecutive passes at 10.0: accept 9.5 as floor.
  - When human input is needed: check for soul files in project directory.

Each dimension has INDEPENDENT scoring logic — scores will differ per document.
  C — Compliance (placeholder check, page-count probe, language markers)
  R — Relevance  (call-text terminology overlap, track/domain markers)
  U — Uniqueness (SotA cues, DOI/refs, specific methodology vs generic terms)
  C — Consortium (partner count, ORCID, CV, gender-balance markers)
  I — Impact     (quantified numbers, open-access, dissemination channels)
  B — Budget     (SEK/EUR amounts, overhead/FTE markers, funding-rate terms)
  L — Leverage   (follow-on funding, scaling language, next-step cues)
  E — Execution  (milestone count, WP refs, risk table, timeline markers)

Usage:
  python crucible_light.py document.pdf --call call_text.txt
  python crucible_light.py document.pdf --call call_text.txt --template template.docx
  python crucible_light.py document.pdf --call call_text.txt --json scores.json
  python crucible_light.py document.pdf --call call_text.txt --iterate --soul-dir ./

License: MIT — WINNIIO AB / Life Atlas
"""

import sys
import re
import json
import argparse
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

try:
    import fitz
    _HAS_FITZ = True
except ImportError:
    _HAS_FITZ = False

try:
    from docx import Document as DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

__version__ = "1.2.0"

# ============================================================
# 8 CALL-SPECIFIC DIMENSIONS (the Light C.R.U.C.I.B.L.E.)
# Each scores 1-10. Markers are weighted keyword groups.
# ============================================================

LIGHT_DIMENSIONS = {
    "compliance": {
        "letter": "C",
        "name": "Compliance",
        "description": "Format, page count, required sections, language, mandatory attachments",
        "weight": 10,
        "checks": [
            {"id": "C1", "name": "Page count within limit", "weight": 20},
            {"id": "C2", "name": "All required sections present", "weight": 25},
            {"id": "C3", "name": "Font/formatting per call spec", "weight": 10},
            {"id": "C4", "name": "Language correct (Swedish/English)", "weight": 15},
            {"id": "C5", "name": "All mandatory attachments listed", "weight": 20},
            {"id": "C6", "name": "No placeholders or [TBD] remaining", "weight": 10},
        ],
        "markers": ["max", "sidor", "bilaga", "obligatorisk", "krav", "format",
                     "font", "mall", "intyg", "undertecknad", "signerad"],
    },
    "relevance": {
        "letter": "R",
        "name": "Relevance",
        "description": "Alignment with call track, non-compensable criterion, domain fit",
        "weight": 20,
        "checks": [
            {"id": "R1", "name": "Explicitly names call track/spår", "weight": 25},
            {"id": "R2", "name": "Problem statement matches call scope", "weight": 25},
            {"id": "R3", "name": "Results strengthen stated capability area", "weight": 20},
            {"id": "R4", "name": "Uses call terminology (not paraphrased)", "weight": 15},
            {"id": "R5", "name": "No scope creep beyond call boundaries", "weight": 15},
        ],
        "markers": ["totalförsvar", "resiliens", "beredskap", "försörjning", "metall",
                     "mineral", "leveranskedja", "förmåga", "spår", "track"],
    },
    "uniqueness": {
        "letter": "U",
        "name": "Uniqueness",
        "description": "Innovation depth, differentiation, state-of-the-art positioning",
        "weight": 15,
        "checks": [
            {"id": "U1", "name": "Clear novelty claim vs state of the art", "weight": 25},
            {"id": "U2", "name": "Named competitors/alternatives acknowledged", "weight": 20},
            {"id": "U3", "name": "Methodology is specific (not generic AI/DT)", "weight": 25},
            {"id": "U4", "name": "Published prior work cited as foundation", "weight": 15},
            {"id": "U5", "name": "Unique combination of approaches justified", "weight": 15},
        ],
        "markers": ["nytt", "unikt", "innovation", "state of the art", "befintlig",
                     "alternativ", "publicerad", "metodik", "validerad", "DOI"],
    },
    "consortium": {
        "letter": "C",
        "name": "Consortium",
        "description": "Partner roles, competence, gender balance, track record",
        "weight": 15,
        "checks": [
            {"id": "K1", "name": "Each partner has clear, non-overlapping role", "weight": 20},
            {"id": "K2", "name": "Key personnel named with credentials", "weight": 20},
            {"id": "K3", "name": "Gender balance addressed quantitatively", "weight": 15},
            {"id": "K4", "name": "Track record / prior collaboration cited", "weight": 20},
            {"id": "K5", "name": "Consortium covers full value chain", "weight": 15},
            {"id": "K6", "name": "Coordinator capacity justified", "weight": 10},
        ],
        "markers": ["konsortium", "partner", "roll", "kompetens", "erfarenhet",
                     "könsbalans", "CV", "ORCID", "nyckelperson", "koordinator"],
    },
    "impact": {
        "letter": "I",
        "name": "Impact",
        "description": "Quantified outcomes, exploitation path, dissemination reach",
        "weight": 15,
        "checks": [
            {"id": "I1", "name": "Quantified impact metrics (not just 'improve')", "weight": 25},
            {"id": "I2", "name": "Exploitation plan names WHO does WHAT after project", "weight": 20},
            {"id": "I3", "name": "Dissemination via credible channels", "weight": 15},
            {"id": "I4", "name": "Number of potential adopters realistic and cited", "weight": 20},
            {"id": "I5", "name": "Open access / replicability plan present", "weight": 20},
        ],
        "markers": ["nyttiggörande", "spridning", "publikation", "open access",
                     "replikerbar", "metodmall", "överförbar", "företag", "bransch"],
    },
    "budget": {
        "letter": "B",
        "name": "Budget",
        "description": "Cost realism, funding rate correctness, overhead limits",
        "weight": 10,
        "checks": [
            {"id": "B1", "name": "Total budget matches sum of partner budgets", "weight": 20},
            {"id": "B2", "name": "Funding rate correct per category + partner type", "weight": 25},
            {"id": "B3", "name": "Overhead ≤30% of personnel (or justified)", "weight": 15},
            {"id": "B4", "name": "Consultant/license costs ≤20% (or motivated)", "weight": 15},
            {"id": "B5", "name": "Person-hours match stated FTE percentages", "weight": 25},
        ],
        "markers": ["budget", "kostnad", "bidrag", "finansiering", "timkostnad",
                     "overhead", "konsult", "licens", "personal", "stödberättigad"],
    },
    "leverage": {
        "letter": "L",
        "name": "Leverage",
        "description": "Next steps, scaling potential, follow-on funding, network effects",
        "weight": 5,
        "checks": [
            {"id": "L1", "name": "Concrete next-step project identified", "weight": 30},
            {"id": "L2", "name": "Named follow-on funding call or investor", "weight": 25},
            {"id": "L3", "name": "Scaling path beyond pilot described", "weight": 25},
            {"id": "L4", "name": "Network/multiplier effect quantified", "weight": 20},
        ],
        "markers": ["nästa steg", "skalning", "uppföljning", "fortsättning",
                     "större ansökan", "följdprojekt", "kommersialisering"],
    },
    "execution": {
        "letter": "E",
        "name": "Execution",
        "description": "Milestones, risks, timeline, dependencies, WP structure",
        "weight": 10,
        "checks": [
            {"id": "E1", "name": "Milestones are measurable (not meetings)", "weight": 20},
            {"id": "E2", "name": "Risk table covers tech + org + external", "weight": 20},
            {"id": "E3", "name": "Timeline is realistic (not front/back-loaded)", "weight": 15},
            {"id": "E4", "name": "WP dependencies mapped and critical path identified", "weight": 20},
            {"id": "E5", "name": "Go/no-go gates present at key decision points", "weight": 15},
            {"id": "E6", "name": "Resource allocation matches WP scope", "weight": 10},
        ],
        "markers": ["milstolpe", "delmål", "leverans", "risk", "tidplan",
                     "arbetspaket", "beroende", "kritisk väg", "go/no-go"],
    },
}


# ============================================================
# TEMPLATE EXPLORATION (Phase 0 — frontloaded)
# ============================================================

@dataclass
class TemplateField:
    name: str
    section: str
    placeholder_text: str
    field_type: str  # "placeholder", "empty_cell", "form_field", "section_header"
    required: bool = True


def explore_template(docx_path: str) -> list:
    """Phase 0: Unpack and map every field/placeholder in template.
    Returns structured inventory of what needs to be filled.
    """
    fields = []
    import zipfile
    import xml.etree.ElementTree as ET

    if not Path(docx_path).exists():
        return fields

    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/document.xml' not in z.namelist():
                return fields
            xml_content = z.read('word/document.xml').decode('utf-8')
    except Exception:
        return fields

    placeholders = re.findall(r'\[([^\]]*(?:Lägg till|insert|TBD|fill|skriv)[^\]]*)\]',
                              xml_content, re.IGNORECASE)
    for p in placeholders:
        fields.append(TemplateField(
            name=p, section="detected", placeholder_text=f"[{p}]",
            field_type="placeholder", required=True))

    form_fields = re.findall(r'w:fldCharType="begin".*?<w:instrText[^>]*>(.*?)</w:instrText>',
                             xml_content, re.DOTALL)
    for ff in form_fields:
        fields.append(TemplateField(
            name=ff.strip(), section="form", placeholder_text="",
            field_type="form_field", required=True))

    return fields


# ============================================================
# SCORING ENGINE
# ============================================================

@dataclass
class DimensionScore:
    key: str
    letter: str
    name: str
    score: float  # 1-10
    max_score: float = 10.0
    checks: dict = field(default_factory=dict)
    gaps: list = field(default_factory=list)
    suggestions: list = field(default_factory=list)


def _count_numbers(text: str) -> int:
    """Count distinct numeric tokens (percentages, large integers, decimals)."""
    return len(re.findall(r'\b\d[\d\s]*[%,.]?\d*\b', text))


def _milestone_count(text: str) -> int:
    """Count milestone/timeline marker occurrences (M1, MS1, WP1, etc.)."""
    return len(re.findall(r'\b(?:m|ms|wp|t)\d+\b', text, re.IGNORECASE))


def _partner_count(text: str) -> int:
    """Estimate number of distinct partner organisations mentioned."""
    # Look for university/company/AB/GmbH/Ltd patterns plus 'partner X'
    hits = re.findall(
        r'\b(?:university|universitet|ab|gmbh|ltd|inc|bv|srl|partner\s+\d)\b',
        text, re.IGNORECASE
    )
    return len(hits)


_CHARS_PER_VIRTUAL_PAGE = 3000


def _extract_text_any(file_path: str) -> tuple[str, int]:
    """Extract full text + page count from PDF, DOCX, MD, or TXT.

    Non-PDF formats approximate page count at ~3000 chars/page.
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        if not _HAS_FITZ:
            print("ERROR: pymupdf required for PDF. Install with: pip install pymupdf")
            sys.exit(1)
        doc = fitz.open(file_path)
        full_text = ""
        page_count = doc.page_count
        for page in doc:
            full_text += page.get_text() + "\n"
        doc.close()
        return full_text, page_count
    elif ext == ".docx":
        if not _HAS_DOCX:
            print("ERROR: python-docx required for .docx. Install with: pip install python-docx")
            sys.exit(1)
        doc = DocxDocument(file_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        text = "\n".join(parts)
        return text, max(1, len(text) // _CHARS_PER_VIRTUAL_PAGE + 1)
    else:
        text = Path(file_path).read_text(encoding="utf-8")
        return text, max(1, len(text) // _CHARS_PER_VIRTUAL_PAGE + 1)


def score_document(pdf_path: str, call_text: str = "") -> dict:
    """Score a document against all 8 Light dimensions with independent logic.

    Each dimension uses its own signals so scores differ per document.
    """
    full_text, page_count = _extract_text_any(pdf_path)

    text_lower = full_text.lower()
    call_lower = call_text.lower() if call_text else ""

    # Shared call-text overlap (used only by Relevance, not shared bonus everywhere)
    call_terms: set = set(re.findall(r'\b\w{4,}\b', call_lower)) if call_lower else set()
    doc_terms: set = set(re.findall(r'\b\w{4,}\b', text_lower))
    call_overlap_ratio = len(call_terms & doc_terms) / max(len(call_terms), 1) if call_terms else 0.0

    scores = {}

    # ------------------------------------------------------------------
    # C — COMPLIANCE
    # Signals: placeholder count (penalise), page count probe, markers
    # ------------------------------------------------------------------
    dim = LIGHT_DIMENSIONS["compliance"]
    placeholder_hits = len(re.findall(
        r'\[\s*(?:tbd|lägg till|insert|skriv|fill in)[^\]]*\]', text_lower))
    marker_found_c = sum(1 for m in dim["markers"] if m in text_lower)
    # Start at 8, penalise placeholders, reward page/format markers
    raw_c = 8.0 - (placeholder_hits * 1.5) + min(2.0, marker_found_c * 0.25)
    # Penalise if doc is suspiciously short (< 5 pages for a full proposal)
    if page_count < 5:
        raw_c -= 1.5

    gap_c = []
    checks_c = {}
    for check in dim["checks"]:
        if check["id"] == "C6":
            passed = placeholder_hits == 0
        elif check["id"] == "C1":
            passed = page_count >= 5
        else:
            passed = marker_found_c >= 2
        checks_c[check["id"]] = {"name": check["name"], "passed": passed, "weight": check["weight"]}
        if not passed:
            gap_c.append(check["name"])

    scores["compliance"] = DimensionScore(
        key="compliance", letter=dim["letter"], name=dim["name"],
        score=round(max(1.0, min(10.0, raw_c)), 1),
        checks=checks_c, gaps=gap_c,
    )

    # ------------------------------------------------------------------
    # R — RELEVANCE
    # Signal: call-text terminology overlap (this dimension's exclusive bonus)
    # ------------------------------------------------------------------
    dim = LIGHT_DIMENSIONS["relevance"]
    marker_found_r = sum(1 for m in dim["markers"] if m in text_lower)
    raw_r = 4.0 + (call_overlap_ratio * 4.0) + min(2.0, marker_found_r * 0.4)

    gap_r = []
    checks_r = {}
    for check in dim["checks"]:
        if check["id"] == "R4":
            passed = call_overlap_ratio > 0.25
        else:
            passed = marker_found_r >= 3
        checks_r[check["id"]] = {"name": check["name"], "passed": passed, "weight": check["weight"]}
        if not passed:
            gap_r.append(check["name"])

    scores["relevance"] = DimensionScore(
        key="relevance", letter=dim["letter"], name=dim["name"],
        score=round(max(1.0, min(10.0, raw_r)), 1),
        checks=checks_r, gaps=gap_r,
    )

    # ------------------------------------------------------------------
    # U — UNIQUENESS
    # Signals: SotA refs, DOI counts, specific vs generic AI/DT language
    # ------------------------------------------------------------------
    dim = LIGHT_DIMENSIONS["uniqueness"]
    marker_found_u = sum(1 for m in dim["markers"] if m in text_lower)
    doi_count = len(re.findall(r'\bdoi\b|10\.\d{4,}/', text_lower))
    generic_hits = len(re.findall(r'\b(?:artificial intelligence|digital twin|machine learning)\b',
                                  text_lower))
    specific_hits = len(re.findall(
        r'\b(?:transformer|lstm|bert|gpt|xgboost|random forest|ontolog|knowledge graph|'
        r'fem|cfd|ansys|comsol|siemens|matlab|openfoam)\b', text_lower, re.IGNORECASE))
    raw_u = 4.0 + min(2.0, doi_count * 0.4) + min(2.0, specific_hits * 0.5) \
            + min(1.5, marker_found_u * 0.3) - min(1.0, generic_hits * 0.2)

    gap_u = []
    checks_u = {}
    for check in dim["checks"]:
        if check["id"] == "U1":
            passed = marker_found_u >= 2
        elif check["id"] == "U2":
            passed = "alternativ" in text_lower or "competitor" in text_lower
        elif check["id"] == "U3":
            passed = specific_hits >= 2
        elif check["id"] == "U4":
            passed = doi_count >= 1
        else:
            passed = marker_found_u >= 3
        checks_u[check["id"]] = {"name": check["name"], "passed": passed, "weight": check["weight"]}
        if not passed:
            gap_u.append(check["name"])

    scores["uniqueness"] = DimensionScore(
        key="uniqueness", letter=dim["letter"], name=dim["name"],
        score=round(max(1.0, min(10.0, raw_u)), 1),
        checks=checks_u, gaps=gap_u,
    )

    # ------------------------------------------------------------------
    # C — CONSORTIUM
    # Signals: partner count, ORCID presence, CV refs, role descriptions
    # ------------------------------------------------------------------
    dim = LIGHT_DIMENSIONS["consortium"]
    marker_found_k = sum(1 for m in dim["markers"] if m in text_lower)
    partners = _partner_count(full_text)
    orcid_hits = len(re.findall(r'orcid|0000-\d{4}-\d{4}-\d{4}', text_lower))
    cv_hits = len(re.findall(r'\bcv\b|\bcurriculum vitae\b', text_lower))
    gender_hits = len(re.findall(r'\b(?:gender|kön|könsbalans|women|men|female|male|kvinn|man)\b',
                                  text_lower))
    raw_k = 3.0 + min(2.0, partners * 0.5) + min(1.5, orcid_hits * 0.75) \
            + min(1.0, cv_hits * 0.5) + min(1.0, gender_hits * 0.5) \
            + min(1.5, marker_found_k * 0.2)

    gap_k = []
    checks_k = {}
    for check in dim["checks"]:
        if check["id"] == "K1":
            passed = partners >= 2
        elif check["id"] == "K2":
            passed = orcid_hits >= 1 or cv_hits >= 2
        elif check["id"] == "K3":
            passed = gender_hits >= 1
        elif check["id"] == "K4":
            passed = "prior" in text_lower or "previous" in text_lower or "tidigare" in text_lower
        else:
            passed = marker_found_k >= 3
        checks_k[check["id"]] = {"name": check["name"], "passed": passed, "weight": check["weight"]}
        if not passed:
            gap_k.append(check["name"])

    scores["consortium"] = DimensionScore(
        key="consortium", letter=dim["letter"], name=dim["name"],
        score=round(max(1.0, min(10.0, raw_k)), 1),
        checks=checks_k, gaps=gap_k,
    )

    # ------------------------------------------------------------------
    # I — IMPACT
    # Signals: quantified numbers, open-access, dissemination markers
    # ------------------------------------------------------------------
    dim = LIGHT_DIMENSIONS["impact"]
    marker_found_i = sum(1 for m in dim["markers"] if m in text_lower)
    number_count = _count_numbers(full_text)
    open_access = 1 if re.search(r'open.?access|öppen tillgång', text_lower) else 0
    dissem_hits = len(re.findall(
        r'\b(?:conference|journal|workshop|linkedin|webinar|konferens|tidskrift)\b',
        text_lower))
    raw_i = 3.0 + min(2.5, number_count * 0.1) + open_access * 1.5 \
            + min(1.5, dissem_hits * 0.4) + min(1.5, marker_found_i * 0.3)

    gap_i = []
    checks_i = {}
    for check in dim["checks"]:
        if check["id"] == "I1":
            passed = number_count >= 5
        elif check["id"] == "I3":
            passed = dissem_hits >= 2
        elif check["id"] == "I5":
            passed = open_access == 1
        else:
            passed = marker_found_i >= 2
        checks_i[check["id"]] = {"name": check["name"], "passed": passed, "weight": check["weight"]}
        if not passed:
            gap_i.append(check["name"])

    scores["impact"] = DimensionScore(
        key="impact", letter=dim["letter"], name=dim["name"],
        score=round(max(1.0, min(10.0, raw_i)), 1),
        checks=checks_i, gaps=gap_i,
    )

    # ------------------------------------------------------------------
    # B — BUDGET
    # Signals: currency amounts (SEK/EUR/kr), overhead/FTE markers
    # ------------------------------------------------------------------
    dim = LIGHT_DIMENSIONS["budget"]
    marker_found_b = sum(1 for m in dim["markers"] if m in text_lower)
    currency_hits = len(re.findall(r'\b(?:sek|eur|kr|euro|kronor)\b|\d[\d\s]*\s*(?:kkr|msek)',
                                    text_lower))
    overhead_hits = len(re.findall(r'\b(?:overhead|oh|indirekta|indirect cost)\b', text_lower))
    fte_hits = len(re.findall(r'\b(?:fte|heltid|deltid|procent av lön|% of salary)\b', text_lower))
    raw_b = 3.0 + min(3.0, currency_hits * 0.5) + min(2.0, overhead_hits * 0.8) \
            + min(1.0, fte_hits * 0.5) + min(1.0, marker_found_b * 0.2)

    gap_b = []
    checks_b = {}
    for check in dim["checks"]:
        if check["id"] == "B1":
            passed = currency_hits >= 3
        elif check["id"] == "B3":
            passed = overhead_hits >= 1
        elif check["id"] == "B5":
            passed = fte_hits >= 1
        else:
            passed = marker_found_b >= 3
        checks_b[check["id"]] = {"name": check["name"], "passed": passed, "weight": check["weight"]}
        if not passed:
            gap_b.append(check["name"])

    scores["budget"] = DimensionScore(
        key="budget", letter=dim["letter"], name=dim["name"],
        score=round(max(1.0, min(10.0, raw_b)), 1),
        checks=checks_b, gaps=gap_b,
    )

    # ------------------------------------------------------------------
    # L — LEVERAGE
    # Signals: follow-on funding mentions, scaling language, next-step cues
    # ------------------------------------------------------------------
    dim = LIGHT_DIMENSIONS["leverage"]
    marker_found_l = sum(1 for m in dim["markers"] if m in text_lower)
    followon_hits = len(re.findall(
        r'\b(?:follow.?on|horizon|vinnova|eurostars|nato|darpa|EIC|eic|seed|series.?[ab])\b',
        text_lower, re.IGNORECASE))
    scale_hits = len(re.findall(
        r'\b(?:scal|skalning|replicat|replikerbar|roll.?out|national|international|global)\b',
        text_lower, re.IGNORECASE))
    raw_l = 2.0 + min(3.0, marker_found_l * 0.8) + min(3.0, followon_hits * 1.0) \
            + min(2.0, scale_hits * 0.5)

    gap_l = []
    checks_l = {}
    for check in dim["checks"]:
        if check["id"] == "L1":
            passed = marker_found_l >= 1
        elif check["id"] == "L2":
            passed = followon_hits >= 1
        elif check["id"] == "L3":
            passed = scale_hits >= 2
        else:
            passed = marker_found_l >= 2
        checks_l[check["id"]] = {"name": check["name"], "passed": passed, "weight": check["weight"]}
        if not passed:
            gap_l.append(check["name"])

    scores["leverage"] = DimensionScore(
        key="leverage", letter=dim["letter"], name=dim["name"],
        score=round(max(1.0, min(10.0, raw_l)), 1),
        checks=checks_l, gaps=gap_l,
    )

    # ------------------------------------------------------------------
    # E — EXECUTION
    # Signals: milestone/WP count, risk table markers, timeline tokens
    # ------------------------------------------------------------------
    dim = LIGHT_DIMENSIONS["execution"]
    marker_found_e = sum(1 for m in dim["markers"] if m in text_lower)
    ms_count = _milestone_count(full_text)
    risk_hits = len(re.findall(r'\b(?:risk|mitigation|åtgärd|sannolikhet|likelihood)\b',
                                text_lower))
    wp_hits = len(re.findall(r'\bwp\s*\d+\b|\barbetspaket\s*\d+\b', text_lower, re.IGNORECASE))
    raw_e = 2.0 + min(3.0, ms_count * 0.4) + min(2.0, risk_hits * 0.4) \
            + min(2.0, wp_hits * 0.4) + min(1.0, marker_found_e * 0.2)

    gap_e = []
    checks_e = {}
    for check in dim["checks"]:
        if check["id"] == "E1":
            passed = ms_count >= 3
        elif check["id"] == "E2":
            passed = risk_hits >= 2
        elif check["id"] == "E4":
            passed = wp_hits >= 2
        else:
            passed = marker_found_e >= 3
        checks_e[check["id"]] = {"name": check["name"], "passed": passed, "weight": check["weight"]}
        if not passed:
            gap_e.append(check["name"])

    scores["execution"] = DimensionScore(
        key="execution", letter=dim["letter"], name=dim["name"],
        score=round(max(1.0, min(10.0, raw_e)), 1),
        checks=checks_e, gaps=gap_e,
    )

    return scores


def composite_score(scores: dict) -> float:
    """Weighted average across all dimensions."""
    total_weight = sum(LIGHT_DIMENSIONS[k]["weight"] for k in scores)
    weighted = sum(scores[k].score * LIGHT_DIMENSIONS[k]["weight"] for k in scores)
    return round(weighted / total_weight, 2) if total_weight else 0


# ============================================================
# SOUL FILE INTEGRATION
# ============================================================

def find_soul_file(search_dir: str = ".") -> dict:
    """Look for a soul file (shadow persona) in the project directory and global shadows.

    Searches (in order):
      - <search_dir>/.claude/agents/shadows/*/
      - <search_dir>/soul_files/
      - <search_dir>/.soul/
      - ~/.claude/agents/shadows/*/   (all shadow subdirs, not just nicolas)

    Returns a dict:
      {
        "path": Path | None,
        "name": str,          # extracted from filename or frontmatter
        "role": str,          # extracted from first heading or "Unknown"
        "context": str,       # first 2000 chars of content
      }
    """
    search = Path(search_dir).resolve()
    shadows_root = Path.home() / ".claude" / "agents" / "shadows"

    candidate_dirs = [
        search / ".claude" / "agents" / "shadows",
        search / "soul_files",
        search / ".soul",
        shadows_root,
    ]

    for base in candidate_dirs:
        if not base.exists():
            continue
        # Depth-1: direct .md/.yaml files in this dir
        direct = list(base.glob("*.md")) + list(base.glob("*.yaml"))
        # Depth-2: subdirectories (each shadow has its own subdir)
        nested: list = []
        for subdir in sorted(base.iterdir()):
            if subdir.is_dir():
                nested += list(subdir.glob("*.md")) + list(subdir.glob("*.yaml"))

        for soul_path in (direct + nested):
            try:
                raw = soul_path.read_text(encoding="utf-8", errors="replace")[:2000]
            except Exception:
                continue

            # Extract name: prefer frontmatter `name:` or stem
            name_match = re.search(r'(?:^|\n)name:\s*([^\n]+)', raw)
            name = name_match.group(1).strip() if name_match else soul_path.stem.title()

            # Extract role: first markdown heading
            role_match = re.search(r'^#\s+(.+)', raw, re.MULTILINE)
            role = role_match.group(1).strip() if role_match else "Unknown"

            return {"path": soul_path, "name": name, "role": role, "context": raw}

    return {"path": None, "name": "Unknown", "role": "Unknown", "context": ""}


def load_soul_context(soul_path: Path) -> str:
    """Load soul file content for human-in-the-loop context."""
    try:
        return soul_path.read_text(encoding='utf-8')[:2000]
    except Exception:
        return ""


def request_human_input(question: str, soul: dict) -> str:
    """Format a human-input request enriched with soul file context.

    When the AI needs a judgment call that requires human expertise
    (e.g. call strategy, consortium fit, budget risk appetite),
    this function returns a formatted prompt surfacing relevant soul context.

    Args:
        question: The specific question the tool needs answered.
        soul: Soul dict from find_soul_file().

    Returns:
        Formatted multi-line string ready to print to terminal.
    """
    lines = [
        "",
        "  [HUMAN INPUT REQUESTED]",
        f"  Expert context: {soul['name']} / {soul['role']}",
        "",
        f"  Q: {question}",
        "",
    ]
    if soul["context"]:
        # Highlight first 3 lines of soul context as framing
        preview = "\n".join(soul["context"].splitlines()[:3])
        lines.append(f"  Soul context preview:")
        for ln in preview.splitlines():
            lines.append(f"    {ln}")
        lines.append("")
    lines.append("  > Enter your answer (blank to skip): ")
    return "\n".join(lines)


# ============================================================
# ITERATION ENGINE
# ============================================================

TARGET_SCORE = 10.0
FLOOR_AFTER_3_PASSES = 9.5
MAX_ITERATIONS = 10


def should_stop(scores: dict, pass_count: int, passes_at_10: int) -> tuple:
    """Determine if iteration should stop.
    Returns (should_stop: bool, reason: str)
    """
    all_scores = [s.score for s in scores.values()]
    min_score = min(all_scores)
    avg_score = sum(all_scores) / len(all_scores)

    if min_score >= TARGET_SCORE:
        return True, f"All dimensions at {TARGET_SCORE} — perfect score"

    if passes_at_10 >= 3 and min_score >= FLOOR_AFTER_3_PASSES:
        return True, f"3 passes at 10, floor {FLOOR_AFTER_3_PASSES} met (min: {min_score})"

    if pass_count >= MAX_ITERATIONS:
        return True, f"Max iterations ({MAX_ITERATIONS}) reached (min: {min_score})"

    return False, f"Pass {pass_count}: min={min_score}, avg={avg_score:.1f} — continuing"


# ============================================================
# REPORTING
# ============================================================

def format_report(scores: dict, pdf_path: str, call_provided: bool,
                  template_fields: list = None) -> str:
    """Format Light CRUCIBLE report."""
    lines = []
    lines.append("=" * 76)
    lines.append(f"  C.R.U.C.I.B.L.E. Light v{__version__}")
    lines.append("  Call-Specific Quick-Score Engine")
    lines.append("=" * 76)
    lines.append(f"  File:       {pdf_path}")
    lines.append(f"  Call text:  {'provided' if call_provided else 'not provided'}")
    lines.append("")

    if template_fields:
        lines.append(f"  TEMPLATE EXPLORATION (Phase 0)")
        lines.append("  " + "-" * 68)
        lines.append(f"  Fields detected: {len(template_fields)}")
        for tf in template_fields[:10]:
            status = "FILLED" if tf.placeholder_text == "" else "EMPTY"
            lines.append(f"    [{status}] {tf.name} ({tf.field_type})")
        if len(template_fields) > 10:
            lines.append(f"    ... and {len(template_fields) - 10} more")
        lines.append("")

    lines.append("  DIMENSIONAL SCORES (1-10)")
    lines.append("  " + "-" * 68)

    for dim_key, dim in LIGHT_DIMENSIONS.items():
        s = scores[dim_key]
        filled = int((s.score - 1) / 9 * 20)
        bar = "#" * filled + "." * (20 - filled)
        grade = "STRONG" if s.score >= 8 else "OK" if s.score >= 6 else "WEAK" if s.score >= 4 else "CRITICAL"
        lines.append(f"  {s.letter} {s.name:<14} {s.score:>4}/10  [{bar}]  {grade}  (w={dim['weight']}%)")

        if s.gaps:
            for gap in s.gaps[:3]:
                lines.append(f"      GAP: {gap}")

    lines.append("")
    comp = composite_score(scores)
    lines.append(f"  COMPOSITE: {comp}/10.0")
    grade = "A+" if comp >= 9.5 else "A" if comp >= 9 else "B" if comp >= 8 else "C" if comp >= 7 else "D" if comp >= 6 else "F"
    lines.append(f"  GRADE: {grade}")
    lines.append("")

    all_gaps = []
    for s in scores.values():
        all_gaps.extend(s.gaps)
    if all_gaps:
        lines.append("  PRIORITY GAPS (fix these first)")
        lines.append("  " + "-" * 68)
        for i, gap in enumerate(all_gaps[:10], 1):
            lines.append(f"  {i}. {gap}")
        lines.append("")

    lines.append("=" * 76)
    return "\n".join(lines)


# ============================================================
# CLI
# ============================================================

def main():
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    parser = argparse.ArgumentParser(
        description="C.R.U.C.I.B.L.E. Light v1.0 — Call-specific quick-score engine",
        epilog="Frontloaded template exploration + 8-dimension scoring. Target: 10/10.",
    )
    parser.add_argument("pdf", help="Path to proposal (PDF, DOCX, MD, or TXT)")
    parser.add_argument("--call", "-c", metavar="PATH",
                        help="Call/topic text file (enables call alignment scoring)")
    parser.add_argument("--template", "-t", metavar="PATH",
                        help="Original .docx template for Phase 0 exploration")
    parser.add_argument("--json", "-j", metavar="PATH",
                        help="Save JSON scores to file")
    parser.add_argument("--iterate", "-i", action="store_true",
                        help="Enable iteration mode (re-score until target)")
    parser.add_argument("--soul-dir", "-s", metavar="PATH", default=".",
                        help="Directory to search for soul files (default: current)")
    parser.add_argument("--verbose", "-v", action="store_true",
                        help="Show detailed check results")

    args = parser.parse_args()

    if not Path(args.pdf).exists():
        print(f"ERROR: File not found: {args.pdf}")
        sys.exit(1)

    call_text = ""
    if args.call:
        call_path = Path(args.call)
        if call_path.exists():
            call_text = call_path.read_text(encoding='utf-8')

    template_fields = []
    if args.template:
        print(f"  Phase 0: Template exploration ({args.template})...")
        template_fields = explore_template(args.template)
        print(f"  Found {len(template_fields)} fields/placeholders")

    print(f"\n  C.R.U.C.I.B.L.E. Light v{__version__}")
    print(f"  Scoring: {args.pdf}")

    scores = score_document(args.pdf, call_text)
    report = format_report(scores, args.pdf, bool(call_text), template_fields)
    print(report)

    if args.iterate:
        soul = find_soul_file(args.soul_dir)
        if soul["path"]:
            print(f"  Soul file found: {soul['path']}  ({soul['name']} / {soul['role']})")
            print(f"  Human context loaded ({len(soul['context'])} chars)")
        else:
            print("  No soul file found — operating autonomously")

        pass_num = 1
        passes_at_10 = 0
        prev_scores: dict = {}

        while True:
            print(f"\n  --- Pass {pass_num} ---")
            current_scores = score_document(args.pdf, call_text)

            if prev_scores:
                print("  Dimension deltas from previous pass:")
                for k, s in current_scores.items():
                    prev = prev_scores[k].score
                    delta = s.score - prev
                    arrow = "+" if delta > 0 else ("" if delta == 0 else "")
                    print(f"    {s.letter} {s.name:<14} {prev:.1f} -> {s.score:.1f}  ({arrow}{delta:+.1f})")

            prev_scores = current_scores
            comp = composite_score(current_scores)
            all_at_10 = all(s.score >= TARGET_SCORE for s in current_scores.values())
            if all_at_10:
                passes_at_10 += 1
                print(f"  passes_at_10 counter: {passes_at_10}")

            stop, reason = should_stop(current_scores, pass_num, passes_at_10)
            print(f"  Stopping condition: {reason}")
            if stop:
                print(f"\n  ITERATION COMPLETE after {pass_num} pass(es). Composite: {comp}")
                break

            # Prompt for manual edit between passes
            input_prompt = request_human_input(
                "Fix the top gaps above and press ENTER when the PDF is updated "
                "(or type 'stop' to exit).",
                soul,
            )
            answer = input(input_prompt).strip().lower()
            if answer == "stop":
                print("  Exiting iteration on user request.")
                break
            pass_num += 1

        scores = prev_scores

    if args.json:
        data = {
            "tool": "CRUCIBLE-Light",
            "version": __version__,
            "file": str(args.pdf),
            "call_provided": bool(call_text),
            "composite": composite_score(scores),
            "dimensions": {
                k: {"letter": s.letter, "name": s.name, "score": s.score,
                     "gaps": s.gaps, "weight": LIGHT_DIMENSIONS[k]["weight"]}
                for k, s in scores.items()
            },
            "template_fields": len(template_fields),
            "iteration_target": TARGET_SCORE,
            "floor_after_3_passes": FLOOR_AFTER_3_PASSES,
        }
        Path(args.json).write_text(json.dumps(data, indent=2), encoding='utf-8')
        print(f"  JSON saved: {args.json}")


if __name__ == "__main__":
    main()
